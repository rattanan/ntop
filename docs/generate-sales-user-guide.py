import os
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "NTOP-Sales-User-Guide-TH.docx"
LOGO = ROOT / "public" / "nt-logo.png"
FONT = "Thonburi"
NAVY = "14324A"
BLUE = "176B87"
TEAL = "2A9D8F"
PALE = "EAF4F4"
LIGHT = "F3F6F8"
GOLD = "E9C46A"
RED = "B33A3A"
GRAY = "5E6B75"
WHITE = "FFFFFF"


def font(run, size=10.5, bold=False, color=NAVY, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(width))
        grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_field(paragraph, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, text, separate, display, end):
        paragraph._p.append(element)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_h(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(p)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    if bold_lead and text.startswith(bold_lead):
        font(p.add_run(bold_lead), bold=True)
        font(p.add_run(text[len(bold_lead):]))
    else:
        font(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        font(p.add_run(item))


def add_steps(doc, items):
    for title, detail in items:
        p = doc.add_paragraph(style="List Number")
        font(p.add_run(title + " — "), bold=True, color=BLUE)
        font(p.add_run(detail))


def callout(doc, label, text, kind="info"):
    colors = {"info": (PALE, TEAL), "warn": ("FFF4D6", "9A6B00"), "risk": ("FBEAEA", RED)}
    fill, accent = colors[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    font(p.add_run(label + "  "), bold=True, color=accent)
    font(p.add_run(text), color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_header(table.rows[0])
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font(p.add_run(value), size=9.2, bold=True, color=WHITE)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            if row_idx % 2:
                shade(cells[idx], LIGHT)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            font(p.add_run(value), size=9.0)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


_page_break_count = 0


def page_break(doc):
    """Reserve separate pages for the cover and orientation page only."""
    global _page_break_count
    _page_break_count += 1
    if _page_break_count <= 2:
        doc.add_page_break()


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, NAVY, 10, 5),
):
    st = styles[name]
    st.font.name = FONT
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    st = styles[name]
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.paragraph_format.left_indent = Inches(0.375)
    st.paragraph_format.first_line_indent = Inches(-0.188)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.25

header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
font(hp.add_run("NTOP  |  คู่มือผู้ใช้งาน Sales"), size=8.5, bold=True, color=GRAY)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(fp.add_run("NTOP Sales Guide  •  "), size=8.5, color=GRAY)
add_field(fp, "PAGE")

# Cover: editorial_cover pattern, compact-reference preset.
doc.add_paragraph().paragraph_format.space_after = Pt(36)
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = p.add_run().add_picture(str(LOGO), width=Inches(1.15))
    doc_pr = picture._inline.docPr
    doc_pr.set("descr", "ตราสัญลักษณ์ NTOP")
    doc_pr.set("title", "NTOP")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(42)
p.paragraph_format.space_after = Pt(8)
font(p.add_run("คู่มือผู้ใช้งาน"), size=14, bold=True, color=TEAL)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
font(p.add_run("NTOP สำหรับ Sales"), size=29, bold=True, color=NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
font(p.add_run("จาก Prospect สู่ Quote ที่ลูกค้ายอมรับ"), size=15, color=BLUE)
callout(doc, "เส้นทางหลัก", "Prospect → Lead → Customer + Opportunity → Solution/Proposal → Quotation/Approval → Accepted", "info")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(64)
font(p.add_run("ฉบับ 1.0  |  อ้างอิงระบบ ณ 26 สิงหาคม 2026"), size=9, color=GRAY)

page_break(doc)
add_h(doc, "วิธีใช้คู่มือฉบับนี้", 1)
add_body(doc, "คู่มือนี้ออกแบบสำหรับ KAM / Sales ที่รับผิดชอบข้อมูลลูกค้า งานติดตาม Opportunity และ Commercial workflow โดยเน้นสิ่งที่ต้องทำในแต่ละวัน จุดตรวจสอบก่อนเปลี่ยนสถานะ และวิธีแก้ปัญหาที่พบบ่อย")
add_table(doc, ["ส่วน", "ใช้เมื่อ"], [
    ("Quick start", "เริ่มวันทำงานและหา record ที่ต้องติดตาม"),
    ("Prospect & Lead", "รับผู้มุ่งหวัง คัดกรอง และแปลงเป็น Customer/Opportunity"),
    ("Opportunity & Pipeline", "ดูแลดีล ปรับ Forecast และจัดการความเสี่ยง"),
    ("Solution, Proposal & Quote", "ประสาน Presales สร้างข้อเสนอ และส่งอนุมัติราคา"),
    ("Troubleshooting & Checklists", "ตรวจสาเหตุเมื่อทำรายการไม่ได้ และทบทวนก่อน Submit"),
], [2100, 7260])
callout(doc, "ขอบเขตสิทธิ์", "เมนู ปุ่ม และข้อมูลที่เห็นขึ้นกับ Role, Organization, Ownership/Assignment และ Workflow responsibility หากคู่มือกล่าวถึงปุ่มที่คุณไม่เห็น ให้ตรวจสิทธิ์และสถานะ record ก่อน", "warn")
add_h(doc, "สารบัญ", 2)
add_bullets(doc, [
    "1. Quick start สำหรับ Sales",
    "2. ค้นหา สร้าง และติดตาม Prospect",
    "3. คัดกรองและ Convert Lead",
    "4. ดูแล Customer และ Opportunity",
    "5. ใช้ Sales Pipeline และ Forecast",
    "6. ประสาน Solution, Survey, BOQ และ Proposal",
    "7. สร้าง Quotation และติดตาม Approval",
    "8. หลังลูกค้ายอมรับ Quote",
    "9. AI Assistance และการคุ้มครองข้อมูล",
    "10. แก้ปัญหาและเช็กลิสต์ส่งงาน",
])

page_break(doc)
add_h(doc, "1. Quick start สำหรับ Sales", 1)
add_h(doc, "1.1 ก่อนเริ่มงาน", 2)
add_steps(doc, [
    ("เข้าสู่ระบบด้วยบัญชีของตนเอง", "ห้ามใช้บัญชีร่วมกัน และหยุดใช้งานทันทีหากชื่อผู้ใช้หรือหน่วยงานไม่ถูกต้อง"),
    ("เปิด Dashboard และกระดิ่งแจ้งเตือน", "ตรวจ Lead ใกล้ SLA, Activity ที่ Overdue, Quote ที่ถูก Return และงานที่ได้รับมอบหมาย"),
    ("เปิด Sales Pipeline", "ทบทวน Commit, Best Case, Expected Close, Next Action และ Risk signals ของดีลที่อยู่ใน scope"),
    ("จัดลำดับงาน", "เริ่มจากงาน Overdue, ดีลใกล้ปิด, งานที่ block Presales/Approval และข้อมูลคุณภาพที่ขาด"),
])
add_h(doc, "1.2 เมนูที่ใช้บ่อย", 2)
add_table(doc, ["เมนู", "งานหลักของ Sales"], [
    ("งานขาย > Prospect", "ค้นหา/สร้างผู้มุ่งหวัง บันทึกการติดตาม และสร้าง Lead เมื่อ Qualified"),
    ("งานขาย > Lead", "Assign, Activity, Qualification และ Convert เป็น Customer + Opportunity"),
    ("งานขาย > Customer", "ดู Customer 360 ผู้ติดต่อ Ownership และประวัติที่เกี่ยวข้อง"),
    ("งานขาย > Opportunity", "บันทึกมูลค่า โอกาสปิด Next Action Requirement และเปลี่ยน Stage"),
    ("Commercial > Sales Pipeline", "ดู Pipeline, Weighted, Forecast mix และ Risk drill-down"),
    ("Commercial > Proposal / Quotation", "สร้างข้อเสนอ สร้าง Quote Version และส่ง Approval"),
    ("งานขาย > กิจกรรม", "บันทึก Call, Email, Meeting, Follow-up หรือ Task"),
], [2600, 6760])
add_h(doc, "1.3 หลักการใช้งานที่ต้องจำ", 2)
add_bullets(doc, [
    "ค้นหาก่อนสร้างเสมอ เพื่อลด Prospect, Lead และ Customer ซ้ำ",
    "ตรวจ Owner, Organization, Status/Stage, Version และ Timeline ก่อนทำคำสั่งสำคัญ",
    "เปลี่ยน Stage/Status ผ่านปุ่ม Workflow ไม่ใช่การแก้ช่องข้อมูลทั่วไป",
    "เมื่อระบบแจ้ง stale version ให้รีเฟรช ตรวจข้อมูลล่าสุด แล้วค่อยทำรายการใหม่",
    "เหตุผลของ Return, Reject, Override หรือ Exception ต้องชัดเจนพอให้ตรวจย้อนหลังได้",
])

page_break(doc)
add_h(doc, "2. ค้นหา สร้าง และติดตาม Prospect", 1)
add_h(doc, "2.1 ค้นหาก่อนสร้าง", 2)
add_body(doc, "เปิด งานขาย > Prospect แล้วค้นหาชื่อบริษัท ชื่อผู้ติดต่อ เลขผู้เสียภาษี เว็บไซต์ อีเมล หรือโทรศัพท์ จากนั้นใช้ตัวกรอง Status, Source, Owner และ Overdue")
callout(doc, "ป้องกันข้อมูลซ้ำ", "ถ้าพบรายการใกล้เคียง ให้เปิดรายละเอียดและตรวจ Owner/Organization ก่อน อย่าสร้างใหม่เพียงเพราะแก้รายการเดิมไม่ได้ เพราะอาจเป็นข้อจำกัดด้านสิทธิ์", "warn")
add_h(doc, "2.2 ลำดับการทำงาน", 2)
add_steps(doc, [
    ("สร้างหรือ Import", "กรอกข้อมูลบริษัท Source, Owner และผู้ติดต่อหลัก; Import ต้อง Preview และแก้ validation error ก่อนยืนยัน"),
    ("มอบหมายและติดต่อ", "ตรวจ Owner แล้วบันทึก Call/Email/Meeting เพื่ออัปเดต Last Contact และ Next Follow-up"),
    ("คัดกรอง", "ตรวจความต้องการ งบประมาณ ผู้มีอำนาจตัดสินใจ ช่วงเวลา และ duplicate candidates"),
    ("ทำให้เป็น QUALIFIED", "กรอกข้อมูลและเหตุผลที่จำเป็นตาม workflow"),
    ("สร้าง Lead", "ทบทวน Transfer Summary และ Qualification Note ก่อนยืนยัน"),
])
add_h(doc, "2.3 สถานะที่พบบ่อย", 2)
add_table(doc, ["สถานะ", "ความหมาย/การกระทำ"], [
    ("NEW / ASSIGNED", "รายการใหม่หรือได้รับมอบหมายแล้ว วางแผน First Contact"),
    ("CONTACTED", "ติดต่อแล้ว บันทึกผลและ Next Follow-up"),
    ("INTERESTED / QUALIFYING", "มีความสนใจและกำลังคัดกรอง"),
    ("QUALIFIED", "พร้อมส่งต่อเป็น Lead"),
    ("CONVERTED", "สร้าง Lead แล้ว ไม่ควรทำซ้ำ"),
], [2200, 7160])

page_break(doc)
add_h(doc, "3. คัดกรองและ Convert Lead", 1)
add_h(doc, "3.1 ดูแล Lead ในแต่ละวัน", 2)
add_bullets(doc, [
    "ใช้ Search, Filter, Sorting, Overdue และ Saved View เพื่อหา Lead ที่ต้องติดตาม",
    "บันทึก Activity เพื่ออัปเดต First-contact SLA, Last Contact และ Next Follow-up",
    "ตรวจ Assignment History และ Status Timeline เมื่อรับงานต่อจากผู้อื่น",
    "Qualification แสดง completeness และ rule score; Manager Override ต้องมีเหตุผล",
])
add_h(doc, "3.2 Convert เป็น Customer + Opportunity", 2)
add_steps(doc, [
    ("ยืนยันว่า Lead ผ่าน Qualification", "สถานะ CONVERTED ใช้ผ่าน Convert command เท่านั้น"),
    ("ตรวจ Duplicate", "เลือก Customer เดิมที่อยู่ใน scope หรือสร้างใหม่; หาก Override duplicate ต้องระบุเหตุผล"),
    ("ตรวจผู้ติดต่อ", "ต้องมีชื่อและช่องทางติดต่อที่จำเป็น และเลือก/สร้าง Contact ให้ถูกคน"),
    ("กรอก Opportunity", "ระบุชื่อดีล Flow, Estimated Value, Expected Close, Probability และ Next Action"),
    ("ยืนยัน Conversion หนึ่งครั้ง", "ระบบสร้าง/เชื่อม Customer, Contact และ Opportunity พร้อม receipt/audit ในธุรกรรมเดียว"),
    ("เปิดผลลัพธ์", "ตรวจ Customer และ Opportunity ที่สร้าง รวมถึง Owner, Organization และ Timeline"),
])
callout(doc, "อย่ากดซ้ำ", "ถ้าหน้าจอยังประมวลผล ให้รอผลก่อน ระบบมี idempotency ป้องกันผลซ้ำ แต่การกดซ้ำทำให้วิเคราะห์ปัญหายากขึ้น", "info")
add_h(doc, "3.3 เมื่อ Convert ไม่สำเร็จ", 2)
add_table(doc, ["อาการ", "ตรวจสอบ"], [
    ("ไม่เห็นปุ่ม Convert", "Qualification, Status, permission, scope และ version"),
    ("เลือก Customer เดิมไม่ได้", "Customer อาจอยู่นอก Organization scope"),
    ("แจ้ง Duplicate", "เลือก Customer เดิม หรือระบุเหตุผล Override ตามสิทธิ์"),
    ("แจ้งข้อมูล Contact", "เติมชื่อและช่องทาง Contact ที่ระบบระบุ"),
    ("Stale version", "รีเฟรชและทบทวนข้อมูลล่าสุด ห้ามเขียนทับการแก้ของผู้อื่น"),
], [2600, 6760])

page_break(doc)
add_h(doc, "4. ดูแล Customer และ Opportunity", 1)
add_h(doc, "4.1 Customer 360", 2)
add_body(doc, "ใช้หน้า Customer เพื่อดูข้อมูลบริษัท ผู้ติดต่อ Ownership, hierarchy, Activity, Lead และ Opportunity ที่เกี่ยวข้อง การเข้าถึงยังคงจำกัดตาม Organization และ record scope")
add_bullets(doc, [
    "ตรวจ Tax ID, ชื่อ, จังหวัด, Segment และผู้ติดต่อก่อนสร้าง Opportunity เพิ่ม",
    "อย่า Merge Customer เองหากไม่มีอำนาจด้าน data governance; Merge ไม่ใช่การลบและมีประวัติถาวร",
    "ข้อมูลต้นทุน Margin และเอกสารอาจถูกซ่อนตาม data classification",
])
add_h(doc, "4.2 สร้างและอัปเดต Opportunity", 2)
add_steps(doc, [
    ("เลือก Customer", "ต้องเป็น Customer ที่เข้าถึงได้และเป็น contracting customer ของดีล"),
    ("ระบุข้อมูลดีล", "ชื่อ Flow, Estimated Value, Probability, Forecast Category และ Expected Close"),
    ("กำหนด Next Action", "เขียนให้เป็นการกระทำที่ชัดเจนและมีวันครบกำหนด"),
    ("เพิ่ม Requirement/Stakeholder", "บันทึกข้อมูลที่ช่วย Stage Gate และการส่งต่อ Presales"),
    ("บันทึกแล้วตรวจ Timeline", "ยืนยัน Owner, Organization และ Version ก่อนเดิน workflow"),
])
add_h(doc, "4.3 Stage path", 2)
add_table(doc, ["Stage", "เป้าหมายก่อนเดินต่อ"], [
    ("QUALIFY", "ยืนยัน fit เบื้องต้น Customer และเจ้าของดีล"),
    ("DISCOVER / NEED ANALYSIS", "Requirement, stakeholder, pain point และ Next Action ชัดเจน"),
    ("SOLUTION", "ประสาน Coverage/Solution/Survey/BOQ ตาม gate"),
    ("PROPOSAL", "Proposal/Quote พร้อมเงื่อนไขเชิงเทคนิคและพาณิชย์"),
    ("NEGOTIATION", "ติดตามข้อเสนอ การอนุมัติ และการตัดสินใจลูกค้า"),
    ("WON / LOST", "บันทึกผล เหตุผล และหลักฐานตาม policy"),
], [2200, 7160])
callout(doc, "Workflow เท่านั้น", "การย้อน Stage, Lost, Cancel, Expire หรือ Reopen ต้องใช้คำสั่ง Transition และระบุเหตุผล ไม่สามารถแก้ Stage แบบลัดจากหน้า Edit", "risk")

page_break(doc)
add_h(doc, "5. ใช้ Sales Pipeline และ Forecast", 1)
add_h(doc, "5.1 อ่าน KPI", 2)
add_table(doc, ["ตัวชี้วัด", "วิธีอ่าน"], [
    ("Open Pipeline", "มูลค่าดีลที่ยังไม่ terminal ภายในช่วงและ scope"),
    ("Weighted Pipeline", "Forecast Amount × Probability คำนวณด้วย Decimal ฝั่ง server"),
    ("Commit", "ดีลที่ทีมขายยืนยันว่ามีความมั่นใจสูงตาม policy"),
    ("Best Case", "ดีลที่อาจปิดได้และรวม Commit ตามสูตรของระบบ"),
    ("Pipeline Coverage", "Eligible Pipeline เทียบ Target; หากไม่มี Target จะแสดง N/A/—"),
], [2500, 6860])
add_h(doc, "5.2 วิธีทบทวน Pipeline", 2)
add_steps(doc, [
    ("เลือกช่วงเวลาและตัวกรอง", "ใช้ Fiscal Period, Owner, Organization, Stage และ Forecast Category"),
    ("ตรวจดีลใกล้ปิด", "Expected Close ต้องสมเหตุผลและมี Next Action ในอนาคต"),
    ("เปิด Risk signals", "แก้ stale stage, no next action, close-date, coverage หรือ approval risk ที่ record ต้นทาง"),
    ("ตรวจ Forecast Amount", "ก่อนมี governed Quote ใช้ Estimated Value; หลังเข้าเงื่อนไขอาจใช้ Primary Quote Version"),
    ("ยืนยันคุณภาพข้อมูล", "อย่าปรับรายงานเพื่อซ่อน warning ให้แก้ Opportunity/Quote ต้นทาง"),
])
callout(doc, "Snapshot", "Forecast snapshot เป็นข้อมูล ณ cutoff และไม่เปลี่ยนย้อนหลังเมื่อ Opportunity ถูกแก้ ใช้ snapshot ล่าสุดสำหรับงานปัจจุบันและ snapshot locked สำหรับการอ้างอิงย้อนหลัง", "info")

page_break(doc)
add_h(doc, "6. ประสาน Solution, Survey, BOQ และ Proposal", 1)
add_h(doc, "6.1 สิ่งที่ Sales ต้องเตรียมให้ Presales", 2)
add_bullets(doc, [
    "Customer, contracting context และผู้ติดต่อที่ถูกต้อง",
    "Requirement, site/location, ข้อจำกัด, timeline และ acceptance expectation",
    "Estimated Value, commercial timing และ Next Action ที่สอดคล้องกับลูกค้า",
    "Owner/ผู้รับผิดชอบและผู้ประสานงานที่ชัดเจน",
])
add_h(doc, "6.2 ติดตาม technical path", 2)
add_table(doc, ["งาน", "หน้าที่ Sales"], [
    ("Coverage / Solution Design", "Request และติดตามผล; Presales/Coverage เป็นผู้ยืนยันข้อมูลเทคนิค"),
    ("Site Survey", "ให้ข้อมูล site/contact และติดตาม assignment; Submit ปัจจุบันเป็น manual snapshot ไม่ใช่ NTSP production success"),
    ("BOQ", "ทบทวนขอบเขต Quantity/Price/Discount; Unit Cost/Margin เห็นตาม permission"),
    ("Technical / Commercial Review", "ตอบข้อซักถาม แก้ requirement และรอผู้มีหน้าที่ review แยกจากผู้ส่ง"),
], [2600, 6760])
add_h(doc, "6.3 สร้าง Proposal", 2)
add_steps(doc, [
    ("เปิด Commercial > Proposal", "เลือก Opportunity ที่อยู่ใน scope"),
    ("ตรวจข้อมูลที่สืบทอด", "Customer, Owner และ Template section ต้องสัมพันธ์กับดีล"),
    ("แก้เนื้อหา", "ตรวจชื่อ คำอธิบาย วันหมดอายุ Tags, scope, assumption และเงื่อนไข"),
    ("บันทึก Version", "การแก้ไขสร้าง immutable version ใหม่; Restore ไม่เขียนทับประวัติ"),
    ("ส่ง Review", "ปฏิบัติตาม Manager/Director review และ maker-checker ที่กำหนด"),
])
callout(doc, "AI Proposal Draft", "AI สร้างได้เพียง editable draft จากข้อมูลที่ได้รับอนุญาต Sales ต้องตรวจข้อเท็จจริง ราคา ขอบเขต และเงื่อนไขก่อนบันทึกหรือส่ง Review", "warn")

page_break(doc)
add_h(doc, "7. สร้าง Quotation และติดตาม Approval", 1)
add_h(doc, "7.1 สร้าง Draft Quote Version", 2)
add_steps(doc, [
    ("สร้าง Quote ใต้ Opportunity", "ตรวจ Customer และ Proposal link ว่าอยู่ใน Opportunity เดียวกัน"),
    ("เพิ่มรายการ", "เลือก Product จำนวน 1–100 รายการ พร้อม Quantity, Unit Price และ Discount"),
    ("ตรวจยอด", "ทบทวน Total, Cost และ Margin ที่ระบบคำนวณด้วย Decimal"),
    ("แก้ Commercial gate", "Floor Price, confirmed cost, Coverage และ Solution gate ต้องครบ"),
    ("ทบทวน Version และ Maker", "ข้อมูลที่ส่งจะถูก snapshot เข้า Approval request"),
    ("Submit Approval", "ส่งเมื่อข้อมูลครบ และติดตามสถานะจาก Quote/Approval queue"),
])
add_h(doc, "7.2 เหตุผลที่อาจถูกยกระดับ", 2)
add_bullets(doc, [
    "Discount มากกว่าเกณฑ์ที่กำหนด (baseline ระบุ >10%)",
    "Gross Margin ต่ำกว่าเกณฑ์ (baseline ระบุ <15%)",
    "Non-standard legal terms, unconfirmed coverage/cost หรือ policy override",
    "Conflict of interest หรือไม่มีผู้อนุมัติที่อยู่ใน authority",
])
add_h(doc, "7.3 เมื่อได้รับ Return หรือ Reject", 2)
add_table(doc, ["ผล", "Sales ต้องทำ"], [
    ("RETURNED", "อ่าน comment แก้ข้อมูลที่ระบุ สร้าง/ใช้ revision ตาม workflow แล้ว Resubmit"),
    ("REJECTED", "ทบทวนเหตุผลและสร้าง Revision ใหม่ก่อน Submit อีกครั้ง"),
    ("APPROVED", "ตรวจว่าทุก mandatory step สำเร็จ แล้วดำเนินการ SENT"),
    ("SENT", "บันทึกผลการส่งและติดตามการตอบรับลูกค้า"),
    ("ACCEPTED", "ยืนยันหลักฐานการยอมรับ ก่อนส่งต่อสร้าง Contract"),
], [2200, 7160])
callout(doc, "Maker-checker", "Maker หรือ Quote editor อนุมัติ mandatory step ของ Quote ตนเองไม่ได้เมื่อ policy กำหนด แม้บัญชีจะมี Role ผู้อนุมัติ", "risk")

page_break(doc)
add_h(doc, "8. หลังลูกค้ายอมรับ Quote", 1)
add_body(doc, "เฉพาะ ACCEPTED Quote Version ที่เข้าถึงได้จึงใช้สร้าง Contract โดย Contract จะสืบทอด Customer, Opportunity, Proposal, Quote และรายการจากฝั่ง server")
add_h(doc, "8.1 สิ่งที่ Sales ต้องส่งต่อ", 2)
add_bullets(doc, [
    "Accepted Quote Version และหลักฐานการยอมรับของลูกค้า",
    "Customer/ผู้ติดต่อและข้อมูล PO หรือเงื่อนไขที่ตกลง",
    "ข้อกำหนด delivery, site, timeline, dependency และ exception",
    "ผู้ประสานงาน Sales/Presales และ Next Action สำหรับ Contract/Order Operations",
])
add_h(doc, "8.2 ข้อจำกัดที่ต้องสื่อสารให้ถูกต้อง", 2)
callout(doc, "Manual handoff", "Service Order ใน workflow ปัจจุบันเป็น DRAFT manual handoff พร้อม immutable snapshot อย่าระบุว่าส่ง NTSP production สำเร็จจนมี acknowledgement/external reference ที่ตรวจสอบได้", "warn")
add_body(doc, "หลังส่งต่อ ให้ติดตาม Activity และสถานะ Contract ตามสิทธิ์ของ Sales แต่ไม่ควรแก้ approval, signature evidence หรือ order state ที่เป็นหน้าที่ของ role อื่น")

add_h(doc, "9. AI Assistance และการคุ้มครองข้อมูล", 1)
add_bullets(doc, [
    "ห้ามใส่ Password, OTP, API key, token, session cookie หรือ credential ลงใน prompt",
    "AI ใช้เฉพาะ visible context และข้อมูลขั้นต่ำตาม capability ไม่ได้เห็นทั้งระบบ",
    "AI Draft/คำอธิบายไม่ใช่ business record จน Sales ตรวจ แก้ และยืนยัน",
    "AI ข้าม workflow, approval, maker-checker, floor price หรือ commercial gate ไม่ได้",
    "เมื่อ AI ไม่พร้อม ให้ใช้ manual Activity/Proposal workflow ได้ งานหลักต้องไม่ถูก block",
])

page_break(doc)
add_h(doc, "10. แก้ปัญหาและเช็กลิสต์ส่งงาน", 1)
add_h(doc, "10.1 ปัญหาที่พบบ่อย", 2)
add_table(doc, ["อาการ", "วิธีตรวจเบื้องต้น"], [
    ("ไม่พบ record", "ตรวจ Organization, Owner/Assignment, filter และ URL; ระบบอาจซ่อน existence เมื่อไม่มีสิทธิ์"),
    ("ไม่เห็นปุ่ม", "ตรวจ Role/permission, Status/Stage, required fields และ workflow responsibility"),
    ("บันทึกไม่ได้", "อ่านข้อความ validation ทีละช่อง และตรวจ Version ก่อน Submit ซ้ำ"),
    ("Stale / Conflict", "รีเฟรช เปิด Timeline ตรวจผู้แก้ล่าสุด แล้วทำรายการจาก Version ปัจจุบัน"),
    ("Quote Submit ไม่ผ่าน", "ตรวจ floor price, cost confirmation, Coverage/Solution gate และรายการสินค้า"),
    ("ยอด Pipeline ไม่เท่า Estimated", "ตรวจ Primary Quote และ policy ของ Forecast Amount"),
    ("AI timeout", "กลับไปใช้ manual workflow และไม่กดส่งซ้ำต่อเนื่อง"),
], [2900, 6460])
add_h(doc, "10.2 เช็กลิสต์ก่อน Submit Proposal/Quote", 2)
add_bullets(doc, [
    "Opportunity, Customer และ Proposal link ถูกต้องและอยู่ใน scope",
    "Requirement, Coverage, Solution, Survey และ BOQ gate ครบตาม Service Category",
    "Product, Quantity, Unit Price, Discount, Floor Price, Cost และ Margin ถูกต้อง",
    "กำลังใช้ Version ล่าสุดและไม่มีข้อมูลที่ยังไม่บันทึก",
    "เหตุผล Exception และเอกสารครบ โดยไม่มี secret/PII เกินจำเป็น",
])
add_h(doc, "10.3 ข้อมูลสำหรับแจ้งปัญหา", 2)
add_bullets(doc, [
    "ชื่อหน้า/เมนู วันเวลา และ timezone (Asia/Bangkok)",
    "Record number/ID ที่มองเห็นได้ตามสิทธิ์",
    "ข้อความ error แบบเต็มและขั้นตอนก่อนเกิดเหตุ",
    "Correlation ID หากหน้าจอแสดง",
    "ภาพหน้าจอที่ปิดบัง PII, ราคา/ต้นทุนลับ และ credential แล้ว",
])
callout(doc, "ห้ามแนบ", "Password, OTP, token, session cookie, private document URL หรือข้อมูลลูกค้าที่ไม่จำเป็นต่อการแก้ปัญหา", "risk")

add_h(doc, "แหล่งอ้างอิงและขอบเขต", 1)
add_body(doc, "คู่มือนี้สรุปจาก approved baseline, workflow/implementation notes, Help Center, navigation และคู่มือรวมของ repository ณ วันที่จัดทำ หาก policy, permission, workflow หรือ route เปลี่ยน ให้ยึด configuration และ source of truth รุ่นล่าสุด")
add_bullets(doc, [
    "docs/product-requirements.md และ docs/system-architecture.md",
    "docs/roles-and-permissions.md, docs/opportunity-workflow.md และ docs/approval-workflow.md",
    "docs/lead-workflow-implementation.md, docs/proposal-quotation-phase1-implementation.md และ docs/sales-forecast-design.md",
    "docs/end-to-end-sales-flow-audit.md และ docs/NTOP-User-Manual-TH.docx",
])

# Core metadata and update fields on open.
doc.core_properties.title = "คู่มือผู้ใช้งาน NTOP สำหรับ Sales"
doc.core_properties.subject = "Sales operating guide"
doc.core_properties.author = "NTOP Product Team"
settings = doc.settings._element
update = OxmlElement("w:updateFields")
update.set(qn("w:val"), "true")
settings.append(update)

if os.environ.get("NTOP_SKIP_SALES_SAVE") != "1":
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)
