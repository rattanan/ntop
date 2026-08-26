import os
import runpy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "NTOP-Admin-User-Guide-TH.docx"
os.environ["NTOP_SKIP_SALES_SAVE"] = "1"
helpers = runpy.run_path(str(ROOT / "docs" / "generate-sales-user-guide.py"))

font = helpers["font"]
set_table_geometry = helpers["set_table_geometry"]
add_field = helpers["add_field"]
add_h = helpers["add_h"]
add_body = helpers["add_body"]
add_bullets = helpers["add_bullets"]
add_steps = helpers["add_steps"]
callout = helpers["callout"]
add_table = helpers["add_table"]
NAVY, BLUE, TEAL, GRAY = helpers["NAVY"], helpers["BLUE"], helpers["TEAL"], helpers["GRAY"]
LOGO, FONT = helpers["LOGO"], helpers["FONT"]


def break_page(doc):
    doc.add_page_break()


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, NAVY, 10, 5),
):
    style = doc.styles[name]
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    style = doc.styles[name]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.188)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

hp = sec.header.paragraphs[0]
font(hp.add_run("NTOP  |  คู่มือผู้ใช้งาน Admin"), size=8.5, bold=True, color=GRAY)
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(fp.add_run("NTOP Admin Guide  •  "), size=8.5, color=GRAY)
add_field(fp, "PAGE")

# Editorial-cover pattern with compact_reference_guide tokens.
doc.add_paragraph().paragraph_format.space_after = Pt(36)
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = p.add_run().add_picture(str(LOGO), width=Inches(1.15))
    picture._inline.docPr.set("descr", "ตราสัญลักษณ์ NTOP")
    picture._inline.docPr.set("title", "NTOP")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(42)
p.paragraph_format.space_after = Pt(8)
font(p.add_run("คู่มือผู้ใช้งาน"), size=14, bold=True, color=TEAL)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
font(p.add_run("NTOP สำหรับ Admin"), size=29, bold=True, color=NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
font(p.add_run("Identity • Organization • Workflow • Audit • AI Governance"), size=14, color=BLUE)
callout(doc, "หลักการสำคัญ", "Admin จัดการ configuration และสิทธิ์ แต่ไม่มี commercial approval โดยอัตโนมัติ ทุก privileged change ต้องผ่าน server authorization และมี audit evidence", "info")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(58)
font(p.add_run("ฉบับ 1.0  |  อ้างอิงระบบ ณ 26 สิงหาคม 2026"), size=9, color=GRAY)

break_page(doc)
add_h(doc, "วิธีใช้คู่มือฉบับนี้", 1)
add_body(doc, "คู่มือนี้สำหรับผู้ดูแลระบบที่ได้รับ capability ตามหน้าที่ ครอบคลุมงานตั้งค่าที่มีผลต่อบัญชี ขอบเขตข้อมูล Approval workflow, AI และหลักฐานตรวจสอบ โดยไม่แทน change approval, deployment runbook หรือ DBA procedure")
add_table(doc, ["ส่วน", "ใช้เมื่อ"], [
    ("Admin quick start", "ตรวจสุขภาพ configuration และงานที่ต้องทบทวน"),
    ("Users & Roles", "สร้าง/แก้บัญชี ปิดใช้งาน และจัดการ role assignment"),
    ("Organization", "จัด hierarchy และผู้อนุมัติ Quotation ในหน่วยงาน"),
    ("Workflow & Authority", "Publish policy, ตั้ง mode, authority และ transition"),
    ("Reference & AI", "จัด Lead rules, service data, provider และ risk rules"),
    ("Audit & Recovery", "ตรวจ Login/Audit และกู้คืนข้อมูลที่ลบแบบ recoverable"),
], [2300, 7060])
callout(doc, "สิทธิ์แยกตาม capability", "การเห็นเมนูไม่ใช่ security control เช่น Users ใช้ user.admin.manage, Organization ใช้ organization.manage และ Audit ใช้ audit.read โดย server ตรวจทุกครั้ง", "warn")
add_h(doc, "สารบัญ", 2)
add_bullets(doc, [
    "1. Admin quick start และหลักปฏิบัติ",
    "2. จัดการบัญชีผู้ใช้และ Role assignments",
    "3. จัดการ Organization และ Quotation approver",
    "4. Approval Control Center และ Workflow",
    "5. Lead management และ reference data",
    "6. AI Provider และ Deal Risk rules",
    "7. Audit, Login History และ Deleted Records",
    "8. Change checklist และการแก้ปัญหา",
])

break_page(doc)
add_h(doc, "1. Admin quick start และหลักปฏิบัติ", 1)
add_h(doc, "1.1 ตรวจสอบก่อนเริ่มงาน", 2)
add_steps(doc, [
    ("ยืนยันบัญชีและหน่วยงาน", "ใช้บัญชีของตนเอง เปิด MFA เมื่อ policy กำหนด และห้ามใช้บัญชีร่วมกัน"),
    ("อ่านคำขอเปลี่ยนแปลง", "ต้องมี owner, เหตุผล, target, effective date และผลกระทบที่คาดไว้"),
    ("ตรวจขอบเขตสิทธิ์", "ยืนยันว่าบัญชีมี capability ที่ตรงกับงาน ไม่อาศัย Role name หรือการเห็นเมนู"),
    ("ตรวจข้อมูลปัจจุบัน", "ทบทวน Version, active/effective period, Organization scope และ dependency ก่อนแก้"),
    ("วางแผนตรวจหลังเปลี่ยน", "กำหนด positive/negative test และ Audit event ที่ต้องพบ"),
])
add_h(doc, "1.2 เมนู Administration", 2)
add_table(doc, ["เมนู", "งานหลัก"], [
    ("ผู้ใช้งาน", "บัญชี สถานะ Role assignments และ API key rotation"),
    ("โครงสร้างหน่วยงาน", "Organization hierarchy, manager/approver assignment และ removal"),
    ("Workflow & Authority", "Global/workflow mode, policy version, authority, transition และ cost confirmation"),
    ("Lead Management", "Campaign/assignment rules และการตั้งค่างาน Lead"),
    ("หมวดหมู่บริการ / Solution reference data", "Reference data และ gate ที่เกี่ยวข้องกับ Solution"),
    ("AI Settings / AI Risk", "Provider configuration และ deterministic Deal Risk rules"),
    ("Audit Log / Deleted Records", "หลักฐาน Login/privileged action และงานกู้คืน"),
], [2900, 6460])
add_h(doc, "1.3 กฎที่ห้ามฝ่าฝืน", 2)
add_bullets(doc, [
    "ห้าม self-grant, self-escalation หรือเปลี่ยน/ถอน privileged role ของตนเอง",
    "ห้ามใส่ password, API key, token หรือ secret ลงในเหตุผล, log, fixture หรือเอกสาร",
    "ห้ามแก้ historical migration หรือ audit evidence เพื่อเปลี่ยนประวัติ",
    "ห้ามเปิด Enforced workflow ก่อนมี Published policy และ approver step ที่ใช้งานได้",
    "ห้ามตีความ Admin role ว่ามีสิทธิ์ approve commercial transaction โดยอัตโนมัติ",
])

add_h(doc, "2. จัดการบัญชีผู้ใช้และ Role assignments", 1)
add_h(doc, "2.1 เปิดรายการและแก้บัญชี", 2)
add_steps(doc, [
    ("เปิด Administration > ผู้ใช้งาน", "ค้นหาบัญชีและตรวจ Status, Role, Organization และเวลาที่แก้ล่าสุด"),
    ("เปิดหน้า Edit", "ใช้หน้ารายละเอียดเฉพาะบัญชี ไม่แก้จากแถวรายการ"),
    ("แก้ข้อมูลที่ได้รับอนุมัติ", "ตรวจ Email/Username, Display name และ Active status"),
    ("จัด Enterprise assignment", "เลือก Role, Scope, Organization และ effective period ให้ตรงคำขอ"),
    ("บันทึกและตรวจ Audit", "ยืนยันผลที่หน้า User และ Audit Log"),
])
callout(doc, "ป้องกัน lockout", "ระบบปฏิเสธการปิดบัญชี เปลี่ยน Role หรือมอบ/ถอน Role ของตนเอง ใช้ผู้ดูแลที่ได้รับมอบหมายอีกคนตามกระบวนการ maker-checker", "risk")
add_h(doc, "2.2 Account lifecycle", 2)
add_table(doc, ["เหตุการณ์", "แนวทาง Admin"], [
    ("Joiner", "ตรวจ sponsor/role owner, ให้ least privilege และส่ง temporary credential ทางช่องทางที่อนุมัติ"),
    ("Mover", "ถอน scope เดิมก่อนเพิ่ม scope ใหม่ และตรวจ overlapping assignment"),
    ("Leaver", "Disable ทันทีและตรวจการ revoke session/token ตาม identity procedure"),
    ("Dormant / review", "ดำเนินการตามรอบ access review และ retention policy"),
], [2300, 7060])
add_h(doc, "2.3 Login และ API key", 2)
add_bullets(doc, [
    "บัญชี Disabled login ไม่ได้ และ session เดิมไม่ผ่าน getSession",
    "Login History แสดง SUCCESS, INVALID_CREDENTIALS หรือ DISABLED แบบ bounded ล่าสุด",
    "การหมุน API key ทำจากหน้า Edit และค่าลับต้องแสดงเพียงครั้งเดียวตาม UI; ห้ามคัดลอกลง ticket/log",
    "Password reset, MFA และ device/session revocation อาจอยู่ใน identity-security milestone แยกจากหน้าจอนี้",
])

add_h(doc, "3. จัดการ Organization และ Quotation approver", 1)
add_h(doc, "3.1 สร้างหรือแก้หน่วยงาน", 2)
add_steps(doc, [
    ("เปิดโครงสร้างหน่วยงาน", "ตรวจ code, name, parent และสถานะของหน่วยงานเป้าหมาย"),
    ("สร้าง/แก้ข้อมูล", "Code ถูก normalize เป็นตัวพิมพ์ใหญ่และต้องไม่ซ้ำ"),
    ("กำหนด Parent", "เลือกเฉพาะ parent ที่ Active; ระบบปฏิเสธ self-parent และ hierarchy cycle"),
    ("บันทึก", "ทบทวนผลต่อ descendant scope และตรวจ Audit event"),
])
add_h(doc, "3.2 มอบหมาย Manager/Approver", 2)
add_steps(doc, [
    ("เลือก User และ Role", "ผู้ใช้ต้อง Active และ Role ต้องตรงกับ active Approval Policy"),
    ("กำหนด scope", "เลือก Organization, effective period, optional segment และ Decimal maximum amount"),
    ("ยืนยัน", "ระบบสร้าง ORG_UNIT role assignment และ approval.decide authority แบบ atomic"),
    ("ทดสอบ", "ตรวจ positive case ใน authority และ negative case ข้าม org/เกินวงเงิน/maker-checker"),
])
add_h(doc, "3.3 ยกเลิกหรือถอดหน่วยงาน", 2)
add_table(doc, ["คำสั่ง", "เงื่อนไขสำคัญ"], [
    ("Revoke approver", "ปิดเฉพาะ role assignment ที่เลือก; shared authority grant ของผู้อื่นยังคงอยู่"),
    ("Remove organization", "เป็น soft delete และต้องไม่มี active child, role assignment หรือ Lead assignment rule"),
    ("Historical links", "Business record เดิมยังอ้างหน่วยงานได้และ Audit history ไม่ถูกลบ"),
], [2700, 6660])

add_h(doc, "4. Approval Control Center และ Workflow", 1)
add_h(doc, "4.1 ลำดับเปิดใช้งานที่ปลอดภัย", 2)
add_steps(doc, [
    ("เริ่มจาก Global Disabled", "Global Disabled มีผลเหนือ workflow รายตัวและไม่มี auto-approve"),
    ("Publish policy version", "กำหนด workflow, tiers/reviewer role, required permission, maker-checker และ effective date"),
    ("สร้าง Authority Grant", "กำหนด Role, permission, Organization/segment, Decimal maximum และ effective period"),
    ("ตรวจ approver coverage", "ต้องมีผู้รับงานที่เข้า scope และ authority ทุก tier ที่จำเป็น"),
    ("เปิด workflow mode", "Enable/Enforce เฉพาะ workflow ที่พร้อม พร้อม reason ที่ชัดเจน"),
    ("เปิด Global mode", "ทำหลังทดสอบ positive/negative และมี rollback/configuration note"),
])
callout(doc, "No auto-approve", "หากไม่มีผู้อนุมัติที่เข้าเงื่อนไข ระบบต้องไม่ auto-approve ให้แก้ assignment/authority/policy อย่างมีหลักฐาน", "risk")
add_h(doc, "4.2 Configuration ที่หน้าเดียวกัน", 2)
add_table(doc, ["ส่วน", "สิ่งที่ต้องตรวจ"], [
    ("Approval workflow", "Amount source, active policy version, mode และ reason"),
    ("Policy tiers", "Tier maximum, reviewer roles, required permission, maker-checker และ effective dates"),
    ("Authority grant", "Role/permission, Organization, segment, maximum amount และช่วงเวลา"),
    ("Opportunity transition", "Policy code, command, from/to stage, required fields และ permission"),
    ("Role assignment", "User, Role, Scope, Organization และ effective dates"),
    ("Confirmed product cost", "Product, Decimal standard cost และ confirmed date"),
], [2700, 6660])
add_h(doc, "4.3 Version และการย้อนกลับ", 2)
add_bullets(doc, [
    "Publish สร้าง policy version ใหม่ ไม่แก้ประวัติเดิม",
    "ใช้ effective dates และ reason เพื่อควบคุมช่วงที่ configuration มีผล",
    "การปิด Authority เป็นการ deactivate ไม่ลบหลักฐานเดิม",
    "ก่อนเปลี่ยน mode ให้บันทึก current version และวิธีคืนค่า configuration เดิม",
])

add_h(doc, "5. Lead management และ reference data", 1)
add_h(doc, "5.1 Lead assignment rules", 2)
add_bullets(doc, [
    "Rule ใช้ priority และเงื่อนไขที่กำหนด โดยรองรับ owner และ organization round robin",
    "ตรวจว่า target owner Active และเป็นสมาชิก Organization ที่ถูกต้อง",
    "หลีกเลี่ยง rule ซ้อนที่ทำให้ assignment ไม่แน่นอน และทดสอบด้วยข้อมูล synthetic",
    "การลบ Organization จะถูก block หากยังมี Lead assignment rule ที่ Active",
])
add_h(doc, "5.2 Service Category และ Solution reference data", 2)
add_table(doc, ["ข้อมูล", "ผลกระทบ"], [
    ("Service Category", "กำหนดว่าต้อง Survey, BOQ หรือ physical-installation gate หรือไม่"),
    ("Product / pricing reference", "มีผลต่อ Quote gate, floor price, confirmed cost และ Margin visibility"),
    ("Solution reference", "มีผลต่อ component, site/survey และ technical workflow options"),
], [2900, 6460])
callout(doc, "Change impact", "Reference data อาจมีผลต่อ record และ workflow ใหม่หลาย module ทดสอบทั้งหน้าสร้าง รายละเอียด การค้นหา และ workflow gate ก่อนเปิดใช้", "warn")

add_h(doc, "6. AI Provider และ Deal Risk rules", 1)
add_h(doc, "6.1 ตั้งค่า AI Provider", 2)
add_steps(doc, [
    ("เปิด Administration > AI Settings", "เฉพาะ Admin ที่ได้รับสิทธิ์เท่านั้น"),
    ("กรอก Endpoint, Model และ Timeout", "ใช้ OpenAI-compatible configuration ที่ได้รับอนุมัติ"),
    ("ใส่ API key", "ระบบเก็บแบบเข้ารหัสและหน้า UI แสดงเพียงว่ามี key ไม่แสดง plaintext"),
    ("บันทึก Version", "การแก้สร้าง configuration version ใหม่และมี Audit"),
    ("Test Connection", "คาดหวัง success หรือ sanitized error เท่านั้น ไม่มี raw provider response/secret"),
    ("Enable เมื่อพร้อม", "ตรวจ AI_CONFIG_MASTER_KEY, fallback และ negative-role test ก่อนเปิด"),
])
add_h(doc, "6.2 Deal Risk rules", 2)
add_bullets(doc, [
    "Rule เป็น deterministic และ versioned; AI ช่วยอธิบายได้แต่ไม่ใช่ trigger authority",
    "กำหนด threshold/severity อย่างมี owner และ effective date",
    "การเปลี่ยน rule ไม่แก้ historical signal; signal ต้องเก็บ rule version และ threshold snapshot",
    "เมื่อ AI Provider ปิด deterministic risk signals ต้องยังแสดงและ workflow หลักต้องทำงานได้",
])
callout(doc, "Secret safety", "ห้ามใส่ provider key, master key, Bearer token หรือ private-key marker ลงใน Audit reason, screenshot, ticket หรือคู่มือ", "risk")

add_h(doc, "7. Audit, Login History และ Deleted Records", 1)
add_h(doc, "7.1 ใช้ Audit อย่างถูกต้อง", 2)
add_bullets(doc, [
    "ค้นจาก actor, action, target, timestamp, result และ correlation ID ตามข้อมูลที่หน้าอนุญาต",
    "Audit/Login History เป็น bounded view; หากต้องสืบสวนกว้างขึ้นให้ใช้กระบวนการ Security/DBA ที่อนุมัติ",
    "Audit เป็น append-only ห้ามแก้หรือลบเพื่อทำให้ผลตรวจเปลี่ยน",
    "ไม่ควรมี password, token, API key, raw prompt/response หรือ PII ที่ไม่จำเป็นในหลักฐาน",
])
add_h(doc, "7.2 Deleted Records และ retention", 2)
add_table(doc, ["Record", "พฤติกรรม"], [
    ("Prospect", "Soft delete พร้อม reason; Admin ค้นหาและ Restore ได้"),
    ("Prospect permanent delete", "ต้องมี dedicated permission และถูก block เมื่อมี Activity/Document/Conversion/History/Receipt/Merge/Audit reference"),
    ("Lead", "ไม่มี delete command; ใช้ Invalid/Archive ตาม lifecycle"),
    ("Opportunity", "ไม่มี delete command; ใช้ Lost/Cancelled/Expired และเก็บรายงานย้อนหลัง"),
    ("Customer", "ไม่มี delete command; ใช้ Inactive/Blacklisted/Closed หรือ governed Merge"),
], [2800, 6560])
callout(doc, "การลบถาวร", "เป็นการกระทำที่กู้คืนยาก ต้องยืนยัน target, permission, dependency และ approved retention basis ก่อนทุกครั้ง", "warn")

add_h(doc, "8. Change checklist และการแก้ปัญหา", 1)
add_h(doc, "8.1 ก่อนเปลี่ยน Configuration", 2)
add_bullets(doc, [
    "มี request owner, business reason, target และ approval ที่ตรวจสอบได้",
    "ตรวจ Version, effective dates, scope, dependency และผลต่อ record เดิม",
    "เตรียม positive, unauthorized, cross-org, maker-checker และ stale-version tests",
    "ไม่ใช้ production secret/PII ใน fixture, screenshot หรือ evidence",
    "กำหนด rollback/configuration restore และผู้รับผิดชอบหลังเปลี่ยน",
])
add_h(doc, "8.2 หลังเปลี่ยน", 2)
add_bullets(doc, [
    "รีเฟรชและยืนยันค่าที่ Active/Effective จริง",
    "ทดสอบด้วยบัญชี role ที่ควรผ่านและควรถูก deny",
    "ตรวจ Audit event, actor, target, result, reason และ correlation",
    "ตรวจว่า core workflow ยังทำงานเมื่อ AI/Integration ปิดหรือไม่พร้อม",
    "บันทึก deployment/configuration note และงานติดตามที่เหลือ",
])
add_h(doc, "8.3 ปัญหาที่พบบ่อย", 2)
add_table(doc, ["อาการ", "ตรวจสอบ"], [
    ("ไม่เห็นเมนู Admin", "capability, effective role assignment, organization และ session ล่าสุด"),
    ("บันทึกไม่ได้", "validation, expectedVersion, self-change protection และ overlapping assignment"),
    ("เปิด Enforced ไม่ได้", "Published policy, approver step, Authority Grant และ Global mode"),
    ("Approver ไม่ได้รับงาน", "Role/permission, Organization/segment, effective dates, maximum amount และ maker-checker"),
    ("ลบ Organization ไม่ได้", "active child, role assignment หรือ Lead assignment rule"),
    ("AI Test Connection ล้มเหลว", "Enabled config, endpoint/model/timeout, configured key และ master key; ห้ามเผย raw secret"),
    ("Restore/Delete ไม่ได้", "permission, expected version และ dependency references"),
], [2900, 6460])
add_body(doc, "เมื่อแจ้งปัญหา ให้แนบชื่อหน้า เวลา Asia/Bangkok, environment, Target ID/Version, error และ Correlation ID โดยปิดบัง PII, commercial-sensitive fields และ secret ทุกครั้ง หาก configuration เปลี่ยนหลังวันที่จัดทำ ให้ยึด source of truth และค่า Active/Effective ล่าสุด")

doc.core_properties.title = "คู่มือผู้ใช้งาน NTOP สำหรับ Admin"
doc.core_properties.subject = "Administration operating guide"
doc.core_properties.author = "NTOP Product Team"
update = OxmlElement("w:updateFields")
update.set(qn("w:val"), "true")
doc.settings._element.append(update)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
